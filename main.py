# main.py (improved version)
import os
import re
import datetime
import uuid
from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify, flash
from docx import Document
import docx
from copy import deepcopy

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Configuration
DOCS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs')
GENERATED_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'generated')
app.config['DOCS_FOLDER'] = DOCS_FOLDER
app.config['GENERATED_FOLDER'] = GENERATED_FOLDER

# Ensure the generated folder exists
os.makedirs(GENERATED_FOLDER, exist_ok=True)

@app.route('/')
def index():
    """Display the home page with directory structure."""
    companies = {}

    # Get the directory structure
    for company_dir in os.listdir(DOCS_FOLDER):
        company_path = os.path.join(DOCS_FOLDER, company_dir)
        if os.path.isdir(company_path):
            files = [f for f in os.listdir(company_path) if f.endswith('.docx')]
            companies[company_dir] = files

    return render_template('index.html', companies=companies)

@app.route('/select/<company>/<filename>')
def select_document(company, filename):
    """Extract placeholders from the selected document and show the form."""
    doc_path = os.path.join(DOCS_FOLDER, company, filename)
    
    if not os.path.exists(doc_path):
        flash(f"Document not found: {filename}", "error")
        return redirect(url_for('index'))

    try:
        # Extract placeholders
        placeholders = extract_placeholders(doc_path)
        return render_template('form.html',
                            company=company,
                            filename=filename,
                            placeholders=placeholders)
    except Exception as e:
        flash(f"Error processing document: {str(e)}", "error")
        return redirect(url_for('index'))

@app.route('/generate', methods=['POST'])
def generate_document():
    """Generate a new document with placeholders replaced."""
    company = request.form['company']
    filename = request.form['filename']
    placeholders = {}

    # Get all placeholder values from the form
    for key, value in request.form.items():
        if key not in ['company', 'filename']:
            placeholders[key] = value

    doc_path = os.path.join(DOCS_FOLDER, company, filename)
    
    if not os.path.exists(doc_path):
        flash(f"Document not found: {filename}", "error")
        return redirect(url_for('index'))

    try:
        # Generate timestamp and unique identifier
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        unique_id = str(uuid.uuid4())[:8]

        # Create new filename with timestamp and unique ID to prevent collisions
        base_name, extension = os.path.splitext(filename)
        new_filename = f"{base_name}_{timestamp}_{unique_id}{extension}"

        generated_path = os.path.join(GENERATED_FOLDER, new_filename)

        # Generate document
        generate_new_document(doc_path, generated_path, placeholders)

        return render_template('download.html', filename=new_filename)
    except Exception as e:
        flash(f"Error generating document: {str(e)}", "error")
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_file(filename):
    """Download the generated file."""
    file_path = os.path.join(GENERATED_FOLDER, filename)
    
    if not os.path.exists(file_path):
        flash(f"Generated file not found: {filename}", "error")
        return redirect(url_for('index'))
        
    return send_file(file_path, as_attachment=True)

def extract_placeholders(doc_path):
    """Extract unique placeholders from a Word document."""
    doc = Document(doc_path)
    placeholders = set()

    # Helper function to process paragraphs
    def process_paragraph(paragraph):
        # Get the full text and identify potential placeholders
        full_text = paragraph.text
        pattern = r'\{\{([^}]+)\}\}'
        
        # Find each potential placeholder
        matches = re.finditer(pattern, full_text)
        for match in matches:
            placeholders.add(match.group(1))
    
    # Process paragraphs in main document
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Process paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                    
    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.footer]:
            if header:
                for paragraph in header.paragraphs:
                    process_paragraph(paragraph)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                process_paragraph(paragraph)

    return sorted(list(placeholders))

def generate_new_document(source_path, target_path, placeholder_values):
    """Generate a new document with placeholders replaced."""
    doc = Document(source_path)
    
    # Process paragraphs in main document
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholder_values)
    
    # Process paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, placeholder_values)
    
    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.footer]:
            if header:
                for paragraph in header.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, placeholder_values)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholders_in_paragraph(paragraph, placeholder_values)
    
    # Save the modified document
    doc.save(target_path)

def replace_placeholders_in_paragraph(paragraph, placeholder_values):
    """
    Replace placeholders in a paragraph, preserving formatting.
    This improved version better handles placeholders split across multiple runs.
    """
    # If paragraph has no runs or text, skip it
    if not paragraph.runs or not paragraph.text:
        return
    
    # Step 1: Build a mapping of placeholder positions in the paragraph text
    full_text = paragraph.text
    placeholder_positions = []
    pattern = r'\{\{([^}]+)\}\}'
    
    for match in re.finditer(pattern, full_text):
        start, end = match.span()
        placeholder_name = match.group(1)
        if placeholder_name in placeholder_values:
            placeholder_positions.append({
                'start': start,
                'end': end,
                'placeholder': match.group(0),  # The full {{placeholder}}
                'name': placeholder_name,       # Just the name inside
                'value': placeholder_values[placeholder_name]
            })
    
    # If no placeholders found, return early
    if not placeholder_positions:
        return
    
    # Step 2: Create a mapping of character positions to runs
    runs = paragraph.runs
    char_to_run_map = []
    current_position = 0
    
    for i, run in enumerate(runs):
        run_length = len(run.text)
        for j in range(run_length):
            char_to_run_map.append({
                'run_index': i,
                'position': j
            })
        current_position += run_length
    
    # Step 3: Process each placeholder
    # Sort in reverse order to avoid position shifts
    for placeholder_info in sorted(placeholder_positions, key=lambda x: x['start'], reverse=True):
        start_pos = placeholder_info['start']
        end_pos = placeholder_info['end']
        replacement = placeholder_info['value']
        
        # Find which runs contain the start and end of the placeholder
        if start_pos >= len(char_to_run_map) or end_pos > len(char_to_run_map):
            # Skip if positions are out of bounds
            continue
            
        start_run_info = char_to_run_map[start_pos]
        end_run_info = char_to_run_map[end_pos - 1] if end_pos > 0 else start_run_info
        
        start_run_index = start_run_info['run_index']
        start_char_pos = start_run_info['position']
        end_run_index = end_run_info['run_index']
        end_char_pos = end_run_info['position'] + 1  # +1 to include the character
        
        # Case 1: Placeholder is entirely within one run
        if start_run_index == end_run_index:
            run = runs[start_run_index]
            before = run.text[:start_char_pos]
            after = run.text[end_char_pos:]
            run.text = before + replacement + after
        
        # Case 2: Placeholder spans multiple runs
        else:
            # Handle first run
            first_run = runs[start_run_index]
            first_run.text = first_run.text[:start_char_pos] + replacement
            
            # Clear intermediate runs
            for i in range(start_run_index + 1, end_run_index):
                runs[i].text = ""
            
            # Handle last run
            last_run = runs[end_run_index]
            last_run.text = last_run.text[end_char_pos:]
    
    # Additional pass to catch any missed placeholders (simpler cases)
    for placeholder_name, replacement in placeholder_values.items():
        placeholder = f"{{{{{placeholder_name}}}}}"
        
        for run in runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

if __name__ == '__main__':
    app.run(debug=True)